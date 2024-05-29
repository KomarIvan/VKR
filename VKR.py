import docx
import openpyxl
import time
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
import spacy
from selenium.common.exceptions import NoSuchElementException

def copy_bibliography_list(file_path):
    doc = docx.Document(file_path)
    bibliography_list = []

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.lower().startswith('библиографический список'):
            start_index = i + 1
            break

    for i in range(start_index, len(doc.paragraphs)):
        bibliography_list.append(doc.paragraphs[i].text)

    return bibliography_list

def create_excel_file(file_path, bibliography_list):
    wb = openpyxl.Workbook()
    sheet = wb.active

    for i, source in enumerate(bibliography_list):
        sheet.cell(row=i + 1, column=1).value = source

    wb.save(file_path)

def get_citation_text(text):

    driver = webdriver.Chrome(service=ChromeService(executable_path='C:/Users/1/Downloads/chromedriver-win64 (1)/chromedriver-win64/chromedriver.exe'))
    driver.get('https://scholar.google.com/schhp?hl=ru&as_sdt=0,5')
    wait = WebDriverWait(driver, 10)

    search_button = driver.find_element(By.XPATH, '//*[@id="gs_hdr_tsb"]')
    search_input = driver.find_element(By.XPATH, '//*[@id="gs_hdr_tsi"]')
    search_input.click()
    search_input.send_keys(text)
    search_button.click()

    time.sleep(5)

    try:
        cite_button = driver.find_element(By.XPATH, '//*[@id="gs_res_ccl_mid"]/div/div/div[3]/a[2]')
        cite_button.click()

        time.sleep(5)

        citation_text = driver.find_element(By.XPATH, '//*[@id="gs_citt"]/table/tbody/tr[1]/td/div').text
    except NoSuchElementException:
        citation_text = "Источник не найден"

    driver.quit()

    return citation_text

def compare_text(source_text, citation_text):
    nlp = spacy.load('ru_core_news_md')
    source_doc = nlp(source_text)
    citation_doc = nlp(citation_text)

    similarity_score = source_doc.similarity(citation_doc)

    return similarity_score

def main():
    docx_file_path = 'C:/Users/1/Downloads/МАГИСТЕРСКАЯ РАБОТА Итог 1 (1).docx'
    excel_file_path = 'C:/Users/1/Downloads/Проверка источников.xlsx'

    bibliography_list = copy_bibliography_list(docx_file_path)
    create_excel_file(excel_file_path, bibliography_list)

    wb = openpyxl.load_workbook(excel_file_path)
    sheet = wb.active

    row_index = 1
    while True:
        source_text = sheet.cell(row=row_index, column=1).value

        if source_text is None:
            break

        citation_text = get_citation_text(source_text)
        sheet.cell(row=row_index, column=2).value = citation_text

        similarity_score = compare_text(source_text, citation_text)
        if citation_text == "Источник не найден":
            sheet.cell(row=row_index, column=3).value = "-"
        else: 
            sheet.cell(row=row_index, column=3).value = similarity_score

        row_index += 1
        wb.save(excel_file_path)
    
    print("Библиографический список проверен успешно.")

if __name__ == "__main__":
    main()